"""
Document loading utilities for different file formats.
Supports PDF, DOCX, and TXT files, with appropriate text extraction.
"""

import os
import re
import io
import tempfile
from typing import List, Dict, Optional, Union, Tuple

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text content as a string
    """
    print(f"Extracting text from PDF: {file_path}")
    text = ""
    
    try:
        try:
            import PyPDF2
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
        except ImportError:
            print("PyPDF2 not installed. Attempting to install...")
            import subprocess
            subprocess.check_call(["pip", "install", "PyPDF2"])
            
            import PyPDF2
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
        
        return text
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {str(e)}")
        return f"[Failed to extract text from {os.path.basename(file_path)}]"

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content as a string
    """
    print(f"Extracting text from DOCX: {file_path}")
    try:
        try:
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except ImportError:
            print("python-docx not installed. Attempting to install...")
            import subprocess
            subprocess.check_call(["pip", "install", "python-docx"])
            
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return f"[Failed to extract text from {os.path.basename(file_path)}]"

def read_text_file(file_path: str) -> str:
    """
    Read content from a plain text file.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        File content as a string
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading text file with latin-1 encoding {file_path}: {str(e)}")
            return ""
    except Exception as e:
        print(f"Error reading text file {file_path}: {str(e)}")
        return ""

def is_scanned_pdf(file_path: str) -> bool:
    """
    Check if a PDF is likely scanned or image-based.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        True if the PDF is likely scanned/image-based, False otherwise
    """
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            num_pages = len(pdf_reader.pages)
            pages_to_check = min(5, num_pages)
            
            total_text = ""
            for i in range(pages_to_check):
                page = pdf_reader.pages[i]
                text = page.extract_text()
                total_text += text
            
            if len(total_text.strip()) < 100 * pages_to_check:
                return True
                
            words = total_text.split()
            single_char_words = sum(1 for word in words if len(word) == 1)
            
            if len(words) > 0 and single_char_words / len(words) > 0.3:
                return True
                
        return False
    except Exception as e:
        print(f"Error checking if PDF is scanned: {str(e)}")
        return True

def load_document(file_path: str) -> Tuple[str, Dict]:
    """
    Load a document and extract its text content based on file type.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Tuple of (text content, metadata)
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # Extract metadata
    metadata = {
        "filename": os.path.basename(file_path),
        "file_path": file_path,
        "file_type": file_ext[1:],  # Remove the dot
        "file_size": os.path.getsize(file_path)
    }
    
    # Extract text based on file type
    if file_ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        text = extract_text_from_docx(file_path)
    elif file_ext == '.txt':
        text = read_text_file(file_path)
    elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp']:
        # Image files are no longer supported without OCR
        text = f"[Image files are not supported: {os.path.basename(file_path)}]"
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")
    
    # Add content length to metadata
    metadata["content_length"] = len(text)
    
    return text, metadata

def split_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
    """
    Split text into smaller chunks for processing.
    
    Args:
        text: Text content to split
        chunk_size: Maximum size of each chunk
        overlap: Number of characters to overlap between chunks
        
    Returns:
        List of text chunks
    """
    chunks = []
    
    if len(text) <= chunk_size:
        return [text]
    
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        
        # Try to find a good breaking point (period, newline, space)
        if end < len(text):
            # Look for a period or newline
            for i in range(end, max(end - 50, start), -1):
                if text[i] in ['.', '\n']:
                    end = i + 1
                    break
            
            # If no good breaking point found, look for a space
            if end == start + chunk_size:
                for i in range(end, max(end - 20, start), -1):
                    if text[i].isspace():
                        end = i + 1
                        break
        
        # Extract the chunk and add to the list
        chunks.append(text[start:end])
        
        # Move start position for next chunk, accounting for overlap
        start = end - overlap
    
    return chunks

def detect_document_type(text: str) -> str:
    """
    Detect if the document is a resume.
    
    Args:
        text: Document text content
        
    Returns:
        "resume" if detected, "unknown" otherwise
    """
    # Look for resume indicators
    resume_indicators = [
        r'\beducation\b', r'\bexperience\b', r'\bskills\b',
        r'\bresume\b', r'\bcv\b', r'\bcareer objective\b',
        r'\bwork history\b', r'\bqualifications\b'
    ]
    
    # Count matches for resume indicators
    resume_score = sum(1 for pattern in resume_indicators if re.search(pattern, text, re.IGNORECASE))
    
    # If the score is too low, we might be uncertain
    if resume_score < 2:
        return "unknown"
    
    return "resume"